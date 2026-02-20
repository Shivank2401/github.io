"""
Generate CV (DOCX + PDF) from portfolio content in docs/index.html.
Run from project root: python generate_cv.py
Output: docs/Shivank_Rajput_CV.docx, docs/Shivank_Rajput_CV.pdf
"""
import os
from pathlib import Path

# Data extracted from docs/index.html
NAME = "Shivank Rajput"
TITLE = "Deputy Manager, Data Analyst and Development"
LOCATION = "Noida, IN"
EMAIL = "rajput.shivank34@gmail.com"
PHONE = "+91 8318875772"
LINKEDIN = "linkedin.com/in/shivank-rajput"
GITHUB = "github.com"

SUMMARY_PARAS = [
    "Results-driven Business Oriented Data Analytics & Strategic Leader with over 6+ years of experience turning chaotic datasets into clear, actionable business roadmaps. My expertise lies in bridging the gap between deep-tier data engineering and high-level executive planning. I don't just build dashboards; I build ecosystems that automate workflows and drive operational excellence.",
    "Throughout my career, I've focused on \"Data with Purpose\"—whether that's slashing inventory defects, optimizing logistics to lower carbon footprints, or mining insights that directly move the needle on sales. I thrive at the intersection of technical precision and strategic vision, ensuring every data point serves a measurable business goal.",
    "I thrive on solving real-world problems, turning ideas into clean, maintainable solutions, and learning through experimentation. You'll find me building data portals, diving into new tech stacks, or exploring what's next in the world of analytics and automation.",
]

KEY_IMPACT = [
    ("Strategic Data Leadership", "Expert at translating complex technical architecture into executive-level insights that inform long-term business strategy."),
    ("Operational Intelligence", "Proven ability to identify bottlenecks in e-commerce and enterprise operations, resulting in a 15% reduction in inventory defect rates through real-time monitoring."),
    ("Growth-Focused Analytics", "Leveraged advanced data mining techniques to identify high-value customer behaviors, leading to a 25% increase in sales conversion rates."),
    ("Scalable Automation", "Specialized in building intelligent automation workflows and Dashboarding ecosystems that replace manual reporting with live, scalable business intelligence."),
    ("Product Architecture", "Deeply experienced in managing the full data lifecycle—from ingestion and transformation to designing the UI/UX for client-side data applications."),
]

EXPERIENCE = [
    {
        "role": "Deputy Manager – Analytics",
        "company": "Eco Soul Home",
        "dates": "Jul 2022 – Present",
        "description": "At Eco Soul Home, I lead analytics efforts on critical business operations, focusing on enhancing data-driven decision-making across multiple platforms. I collaborate closely with C-suite executives and stakeholders to ensure that data solutions meet both operational needs and strategic objectives.",
        "bullets": [
            "<b>Data Source & Integration</b>: Collected and integrated data from multiple sources including databases, APIs, Azure Blob Storage, and operational systems.",
            "<b>Data Engineering & Pipeline Development</b>: Built automated ETL/ELT pipelines using Python, SQL, and Azure Data Factory to transform raw data into structured analytical datasets.",
            "<b>Data Modeling & Processing</b>: Designed scalable data models and optimized large datasets for efficient querying, reporting, and analytics.",
            "<b>Analytics, Reporting & Visualization</b>: Developed interactive dashboards and KPI reports using Power BI, Excel, and Python to support business decision-making.",
            "<b>Operational Intelligence & Automation</b>: Automated repetitive workflows, inventory tracking, and operational reports to improve efficiency and accuracy.",
            "<b>Cost Optimization & Business Impact</b>: Identified cost-saving opportunities through data analysis, reducing operational defects by 15% and improving logistics efficiency by 20%.",
            "<b>Strategic Insights & Decision Support</b>: Translated complex data into actionable insights for leadership, driving conversion growth and process improvements.",
            "<b>Stakeholder Collaboration</b>: Worked closely with cross-functional teams including product, operations, and business stakeholders to align analytics with organizational goals.",
            "<b>Team Leadership & Mentorship</b>: Led analytics initiatives, guided junior analysts, and ensured timely delivery of high-impact data projects.",
        ],
    },
    {
        "role": "Business Analyst",
        "company": "Data Flow Group",
        "dates": "Jan 2022 – Jun 2022",
        "description": "As a Business Analyst at Data Flow Group, I was responsible for analyzing PSV data and providing detailed insights for executive decision-making. My role involved conducting rigorous data audits, reporting, Monthly Board Decks and collaborating with HR Department on various analysis.",
        "bullets": [
            "Analyzed revenue inflow, PSV, EBITDA, book revenues, etc. trends to provide diagnostic insights for executive decision-making",
            "Creating Monthly Board decks, Operational KPIs, Audit deck, Process Flow Charts etc.",
            "Conducted rigorous data audits at case and check levels, ensuring 99.9% data integrity for financial reporting",
            "Collaborated with HR leadership to analyze attrition data, identifying key drivers of employee turnover",
        ],
    },
    {
        "role": "Business Analyst & Development Executive",
        "company": "Zoylee Web Services",
        "dates": "Nov 2019 – Jan 2022",
        "description": "At Zoylee Web Services, I supported business growth through data-driven strategies and CRM optimization. I focused on geo-spatial data mining and customer engagement strategies while learning and growing in a fast-paced analytics environment.",
        "bullets": [
            "Increased sales conversion rates by 25% through geo-spatial data mining and city-level market mapping",
            "Improved customer retention by 10% by leveraging CRM data to personalize engagement strategies",
            "Identified critical business threats through SWOT analysis and proposed data-backed mitigation strategies",
        ],
    },
    {
        "role": "Business Development Executive",
        "company": "InnovationM Technology Solution, Noida",
        "dates": "Jan 2019 – Apr 2019 (+ 3 months Internship)",
        "description": "At InnovationM Technology Solution, I drove business development through prospecting, client relationship management, and staff augmentation for US-based and domestic clients, with focus on Contract-to-Permanent (CTP) and Contract-to-Hire (CTH) recruitment.",
        "bullets": [
            "Prospecting and identifying new sales leads, while also nurturing and maintaining positive relationships with current clientele.",
            "Conducting staff augmentation for both US-based and domestic clients, focusing on Contract-to-Permanent (CTC) or Contract-to-Hire (CTH) roles and recruitment efforts.",
            "Identifying potential new markets and leads, initiating contact with prospective clients through email or phone to build rapport and schedule meetings.",
            "Managing databases for Cold Calling and email marketing campaigns to generate prospective leads.",
        ],
    },
]

DOMAIN_EXPERTISE = [
    "Amazon/Flipkart/Walmart APIs", "Retail & E-commerce Ops", "CXO Board Deck Creation",
    "Project Management", "Ops Workflow Management", "Team Management & Training", "Cross-functional Collaboration",
]
TECHNICAL_ANALYTICS = [
    "Data Analysis & Mining", "MIS Reporting Architecture", "Reporting Automation",
    "Process Optimization", "Primary & Secondary Research", "Insight Generation",
]
TECH_STACK = [
    ("Programming & Development", "Python, VS Code, Office 365, Power Automate"),
    ("Databases & Infrastructure", "Azure Blob Storage, Azure App Services, MongoDB, Qdrant, phpMyAdmin"),
    ("Data Visualization, Reporting & Design", "Power BI, PowerPoint, Prezi, Figma, Canva, Adobe Illustrator"),
    ("Productivity & Office Tools", "MS Office Suite (Word, Excel, PowerPoint), Smartsheet, Google Sheets"),
]

PROJECTS = [
    ("Datahive", "Centralized data warehouse using Azure and Power BI, establishing a Single Source of Truth."),
    ("HRMS Suite", "Admin & Employee portals for performance tracking and attrition analysis."),
    ("Ops Trackers", "Project, Query, and Asset Trackers — digitized workflows with real-time visibility."),
    ("Demand Planner Workflow", "Hierarchy-based approval system for demand planning with data validation before ERP ingestion."),
    ("Internal Operations Suite", "Digitized manual operational logs into interactive applications using Power Apps and Power Automate."),
    ("HRMS Portal Architecture", "Designed the logic and data structure for Admin and Employee portals to automate performance tracking and attrition analysis."),
    ("Inventory & Ops Dashboard", "Real-time inventory operations dashboards tracking SLA and aging for defect reduction."),
    ("Sales Analytics Dashboard", "P&L analysis, customer buying patterns, and returns analysis for strategic decisions."),
    ("Board Decks & CXO Reports", "Monthly board decks and executive reports translating operational metrics into growth plans."),
    ("Strategy Presentations", "Data-backed strategy and business review presentations for leadership and stakeholders."),
    ("ETL & Reporting Automation", "Python and Power Automate pipelines ingesting Amazon/Walmart data, reducing manual reporting by 90%."),
    ("Carbon Footprint Tracker", "Tracking system for emissions monitoring and validation of sustainability targets."),
]

EDUCATION = [
    ("Master of Business Administration-Honors (Marketing & HR)", "PSIT, Kanpur | 2017 – 2019"),
    ("Bachelor of Technology (Mechanical Engineering)", "K.N.G.D. Modi Engineering College, Ghaziabad | 2011 – 2015"),
]
CERTIFICATIONS = [
    "Fundamentals of Operation Management | Alison",
    "Strategic Management | Alison",
    "Research Abstract Published under C.S.J.M University",
]
LANGUAGES = [("English", 90), ("Hindi", 100), ("Japanese (Basic)", 10)]

RESEARCH_ABSTRACT = [
    "Role of Engineering in Sustainable Agriculture Development",
    "Impact on Climate Change on Agriculture & Food Security",
    "Climate Change: Natural Disasters & Migration",
    "Renewable Energy & Agriculture Development",
]
MBA_RESEARCH = [
    "To Evaluate the impact of integrated Mobility Solution (IMS) & Digital Transformation on Users",
    "A Study on Digitalization Changing Phase & Growth of IT Industry.",
]


def create_docx(out_path: Path):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 51, 102)
        p_e = p._p
        pPr = p_e.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "003366")
        pbdr.append(bottom)
        pPr.append(pbdr)

    # Header
    h = doc.add_heading(NAME, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    sub.add_run(TITLE + "\n")
    sub.add_run(f"{LOCATION} | {PHONE} | {EMAIL} | {LINKEDIN}")
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].bold = True
    for r in sub.runs[1:]:
        r.font.size = Pt(10)

    # Summary
    section_header("Professional Summary")
    for para in SUMMARY_PARAS:
        doc.add_paragraph(para)
    p_impact = doc.add_paragraph()
    p_impact.paragraph_format.space_before = Pt(8)
    r = p_impact.add_run("Key Impact & Expertise")
    r.bold = True
    r.font.size = Pt(11)
    for title, desc in KEY_IMPACT:
        doc.add_paragraph(f"• {title}: {desc}", style="List Bullet")

    # Work Experience
    section_header("Work Experience")
    for job in EXPERIENCE:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{job['company']}  |  {job['role']}")
        r1.bold = True
        p.add_run(f"\t{job['dates']}")
        doc.add_paragraph(job["description"]).paragraph_format.space_after = Pt(4)
        for b in job["bullets"]:
            doc.add_paragraph(b, style="List Bullet").paragraph_format.space_after = Pt(2)

    # Skills
    section_header("Skills & Tools")
    p_skill = doc.add_paragraph()
    p_skill.add_run("Domain Expertise: ").bold = True
    p_skill.add_run(", ".join(DOMAIN_EXPERTISE) + "\n")
    p_skill.add_run("Technical Analytics: ").bold = True
    p_skill.add_run(", ".join(TECHNICAL_ANALYTICS) + "\n")
    for cat, items in TECH_STACK:
        p_skill.add_run(f"{cat}: ").bold = True
        p_skill.add_run(items + "\n")

    # Selected Projects
    section_header("Selected Projects")
    for title, desc in PROJECTS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.add_run(title).bold = True
        doc.add_paragraph(desc, style="List Bullet")

    # Education
    section_header("Education")
    for degree, detail in EDUCATION:
        p = doc.add_paragraph()
        p.add_run(degree).bold = True
        doc.add_paragraph(detail)

    # Certifications
    section_header("Certifications")
    for c in CERTIFICATIONS:
        doc.add_paragraph("• " + c)

    # Languages
    section_header("Languages")
    doc.add_paragraph(", ".join(f"{lang} ({pct}%)" for lang, pct in LANGUAGES))

    # Research
    section_header("Research Works")
    p_abs = doc.add_paragraph()
    p_abs.add_run("Abstract Published under CSJMU, Kanpur University").bold = True
    for r in RESEARCH_ABSTRACT:
        doc.add_paragraph("• " + r)
    p_mba = doc.add_paragraph()
    p_mba.add_run("MBA Research Reports").bold = True
    for r in MBA_RESEARCH:
        doc.add_paragraph("• " + r)

    doc.save(str(out_path))
    print("Created:", out_path)


def create_pdf(out_path: Path):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        name="CVTitle",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=4,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#000000"),
    )
    sub_style = ParagraphStyle(
        name="CVSub",
        parent=styles["Normal"],
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=14,
    )
    head_style = ParagraphStyle(
        name="SectionHead",
        parent=styles["Heading2"],
        fontSize=11,
        spaceBefore=12,
        spaceAfter=6,
        textColor=colors.HexColor("#003366"),
        leftIndent=0,
    )
    body = styles["Normal"]
    body.fontSize = 10
    body.spaceAfter = 4
    bullet_style = ParagraphStyle(
        name="Bullet",
        parent=body,
        leftIndent=20,
        bulletIndent=8,
        spaceAfter=3,
    )

    story = []

    story.append(Paragraph(NAME, title_style))
    story.append(Paragraph(TITLE, sub_style))
    story.append(Paragraph(f"{LOCATION} | {PHONE} | {EMAIL} | {LINKEDIN}", sub_style))
    story.append(Spacer(1, 6))

    story.append(Paragraph("PROFESSIONAL SUMMARY", head_style))
    for para in SUMMARY_PARAS:
        story.append(Paragraph(para.replace("&", "&amp;"), body))
    story.append(Paragraph("<b>Key Impact & Expertise</b>", body))
    for title, desc in KEY_IMPACT:
        story.append(Paragraph(f"• <b>{title}:</b> {desc.replace('&', '&amp;')}", bullet_style))

    story.append(Paragraph("WORK EXPERIENCE", head_style))
    for job in EXPERIENCE:
        story.append(Paragraph(
            f"<b>{job['company']}  |  {job['role']}</b>  {job['dates']}".replace("&", "&amp;"),
            body,
        ))
        story.append(Paragraph(job["description"].replace("&", "&amp;"), body))
        for b in job["bullets"]:
            story.append(Paragraph("• " + b.replace("&", "&amp;"), bullet_style))
        story.append(Spacer(1, 4))

    story.append(Paragraph("SKILLS & TOOLS", head_style))
    story.append(Paragraph("<b>Domain Expertise:</b> " + ", ".join(DOMAIN_EXPERTISE), body))
    story.append(Paragraph("<b>Technical Analytics:</b> " + ", ".join(TECHNICAL_ANALYTICS), body))
    for cat, items in TECH_STACK:
        story.append(Paragraph(f"<b>{cat}:</b> {items}", body))

    story.append(Paragraph("SELECTED PROJECTS", head_style))
    for title, desc in PROJECTS:
        story.append(Paragraph(f"<b>{title}</b>", body))
        story.append(Paragraph("• " + desc.replace("&", "&amp;"), bullet_style))

    story.append(Paragraph("EDUCATION", head_style))
    for degree, detail in EDUCATION:
        story.append(Paragraph(f"<b>{degree}</b>", body))
        story.append(Paragraph(detail, body))
    story.append(Paragraph("CERTIFICATIONS", head_style))
    for c in CERTIFICATIONS:
        story.append(Paragraph("• " + c, body))
    story.append(Paragraph("LANGUAGES", head_style))
    story.append(Paragraph(", ".join(f"{lang} ({pct}%)" for lang, pct in LANGUAGES), body))
    story.append(Paragraph("RESEARCH WORKS", head_style))
    story.append(Paragraph("<b>Abstract Published under CSJMU, Kanpur University</b>", body))
    for r in RESEARCH_ABSTRACT:
        story.append(Paragraph("• " + r, body))
    story.append(Paragraph("<b>MBA Research Reports</b>", body))
    for r in MBA_RESEARCH:
        story.append(Paragraph("• " + r, body))

    doc.build(story)
    print("Created:", out_path)


def main():
    root = Path(__file__).resolve().parent
    docs_dir = root / "docs"
    docs_dir.mkdir(exist_ok=True)
    base = docs_dir / "Shivank_Rajput_CV"
    create_docx(Path(str(base) + ".docx"))
    try:
        create_pdf(Path(str(base) + ".pdf"))
    except ImportError:
        print("PDF skipped: install reportlab with 'pip install reportlab' to generate PDF.")
    print("Done. CV files are in docs/")


if __name__ == "__main__":
    main()
