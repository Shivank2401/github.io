from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_ats_cv():
    document = Document()

    # --- Styles & Fonts ---
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'  # Arial is extremely safe for ATS
    font.size = Pt(10.5)

    # Helper to add a clean section border
    def add_section_header(text):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0) # Black for high contrast
        
        # Add bottom border for the header
        p_element = p._p
        pPr = p_element.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pbdr.append(bottom)
        pPr.append(pbdr)

    # --- 1. Header Information ---
    header = document.add_heading('SHIVANK RAJPUT', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subheader / Contact
    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(12)
    contact_run = contact.add_run('Noida, IN | +91 8318875772 | rajput.shivank34@gmail.com | linkedin.com/in/shivank-rajput') 
    contact_run.font.size = Pt(10)

    # --- 2. Professional Summary ---
    add_section_header('Professional Summary')
    summary_text = (
        "Strategic Analytics & Operations Leader with 5+ years of experience in Data Engineering, Business Intelligence, "
        "and Product Management. Expert in architecting internal data portals (HRMS, Datahive) and automating ETL workflows "
        "using Python and Power Automate. Proven track record of driving 15-25% efficiency gains in Inventory Management and "
        "Sales Operations for high-volume E-commerce platforms (Amazon/Flipkart). Adept at translating complex datasets into "
        "actionable CXO-level strategies."
    )
    document.add_paragraph(summary_text)

    # --- 3. Technical Skills (Linear Layout for ATS) ---
    add_section_header('Technical Skills & Tools')
    
    skills_format = document.add_paragraph()
    skills_format.paragraph_format.line_spacing = 1.15
    
    # Helper to add skill categories
    def add_skill_category(category, items):
        run_cat = skills_format.add_run(f"{category}: ")
        run_cat.bold = True
        skills_format.add_run(f"{items}\n")

    add_skill_category("Data & ETL", "Python (Pandas, NumPy), SQL, Azure, MongoDB, Power Query, ETL Pipelines")
    add_skill_category("Visualization & BI", "Power BI (DAX), Tableau, Google Data Studio, Advanced Excel, Prezi")
    add_skill_category("Automation & Dev", "Power Automate, VS Code, Git, SharePoint, REST APIs")
    add_skill_category("Domain Knowledge", "Amazon Seller Central, Flipkart/Walmart Ops, Inventory Management, CRM")
    add_skill_category("Design & Management", "Figma, Canva, Project Management (Agile), Team Leadership")

    # --- 4. Professional Experience ---
    add_section_header('Professional Experience')

    # Job 1
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    line1 = p.add_run('Eco Soul Home | Deputy Manager - Analytics')
    line1.bold = True
    p.add_run('\t\t\t\t\tJul 2022 – Present') # Tab for right align date
    
    bullets_1 = [
        "Spearheaded the development of 'Datahive', a centralized data warehouse using Azure and Power BI, replacing fragmented reporting and establishing a 'Single Source of Truth'.",
        "Architected automated ETL pipelines using Python and Power Automate to ingest raw sales data from Amazon/Walmart, reducing manual reporting time by 90%.",
        "Developed real-time Inventory Operations Dashboards tracking SLA and aging, resulting in a 15% reduction in defect rates.",
        "Designed and implemented a Carbon Footprint Tracking System, successfully monitoring and validating a 20% reduction in emissions.",
        "Partnered with C-suite executives to deliver monthly Board Decks, translating operational metrics into strategic growth plans."
    ]
    for b in bullets_1:
        bp = document.add_paragraph(b, style='List Bullet')
        bp.paragraph_format.space_after = Pt(2)

    # Job 2
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    line2 = p.add_run('Data Flow Group | Business Analyst')
    line2.bold = True
    p.add_run('\t\t\t\t\t\t\tJan 2022 – Jun 2022')
    
    bullets_2 = [
        "Analyzed revenue inflow and PSV EBITDA trends to provide diagnostic insights for executive decision-making.",
        "Conducted rigorous data audits at case and check levels, ensuring 99.9% data integrity for financial reporting.",
        "Collaborated with HR leadership to analyze attrition data, identifying key drivers of employee turnover."
    ]
    for b in bullets_2:
        bp = document.add_paragraph(b, style='List Bullet')
        bp.paragraph_format.space_after = Pt(2)

    # Job 3
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    line3 = p.add_run('Zoylee Web Services | Business Analyst & Development Executive')
    line3.bold = True
    p.add_run('\t\t\tNov 2019 – Jan 2022')
    
    bullets_3 = [
        "Increased sales conversion rates by 25% through geo-spatial data mining and city-level market mapping.",
        "Improved customer retention by 10% by leveraging CRM data to personalize engagement strategies.",
        "Identified critical business threats through SWOT analysis and proposed data-backed mitigation strategies."
    ]
    for b in bullets_3:
        bp = document.add_paragraph(b, style='List Bullet')
        bp.paragraph_format.space_after = Pt(2)

    # --- 5. Key Projects & Portals ---
    add_section_header('Key Projects & Product Management')
    
    # Project 1
    proj_p = document.add_paragraph()
    proj_p.paragraph_format.space_before = Pt(4)
    run_proj = proj_p.add_run("Internal Operations Suite (Project, Query, & Asset Trackers)")
    run_proj.bold = True
    proj_desc = document.add_paragraph("Digitized manual operational logs into interactive applications using Power Apps and Power Automate, providing real-time visibility into asset allocation and project status.", style='List Bullet')

    # Project 2
    proj_p2 = document.add_paragraph()
    proj_p2.paragraph_format.space_before = Pt(2)
    run_proj2 = proj_p2.add_run("HRMS Portal Architecture")
    run_proj2.bold = True
    proj_desc2 = document.add_paragraph("Designed the logic and data structure for Admin and Employee portals to automate performance tracking and attrition analysis.", style='List Bullet')

    # Project 3
    proj_p3 = document.add_paragraph()
    proj_p3.paragraph_format.space_before = Pt(2)
    run_proj3 = proj_p3.add_run("Demand Planner Workflow Tool")
    run_proj3.bold = True
    proj_desc3 = document.add_paragraph("Built a hierarchy-based approval system for demand planning, ensuring data validation before ingestion into ERP systems.", style='List Bullet')

    # --- 6. Education ---
    add_section_header('Education')
    
    edu = document.add_paragraph()
    edu.paragraph_format.space_after = Pt(0)
    edu_run = edu.add_run("Master of Business Administration (Marketing & HR)")
    edu_run.bold = True
    document.add_paragraph("PSIT, Kanpur | 2017 – 2019")

    edu2 = document.add_paragraph()
    edu2.paragraph_format.space_before = Pt(4)
    edu2.paragraph_format.space_after = Pt(0)
    edu2_run = edu2.add_run("Bachelor of Technology (Mechanical Engineering)")
    edu2_run.bold = True
    document.add_paragraph("K.N.G.D. Modi Engineering College, Ghaziabad | 2011 – 2015")

    # --- 7. Certifications ---
    add_section_header('Certifications')
    cert_p = document.add_paragraph()
    cert_p.add_run("• Fundamentals of Operation Management | Alison\n")
    cert_p.add_run("• Strategic Management | Alison\n")
    cert_p.add_run("• Research Abstract Published under C.S.J.M University")

    document.save('Shivank_Rajput_ATS_CV.docx')

create_ats_cv()